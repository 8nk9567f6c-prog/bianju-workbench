"""Markdown to Word converter for 编剧工作台.
Usage: python md2docx.py <input.md> <output.docx>
"""

import sys
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, lines, start_idx):
    """Parse a markdown table and add it to the document."""
    header_line = lines[start_idx]
    sep_line = lines[start_idx + 1]

    headers = [h.strip() for h in header_line.split("|")[1:-1]]
    aligns = []
    for part in sep_line.split("|")[1:-1]:
        part = part.strip()
        if part.startswith(":") and part.endswith(":"):
            aligns.append(WD_ALIGN_PARAGRAPH.CENTER)
        elif part.endswith(":"):
            aligns.append(WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            aligns.append(WD_ALIGN_PARAGRAPH.LEFT)

    table = doc.add_table(rows=1, cols=len(headers), style="Table Grid")
    table.autofit = True

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        p.alignment = aligns[i] if i < len(aligns) else WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, "D9E2F3")

    # Data rows
    row_idx = start_idx + 2
    while row_idx < len(lines) and lines[row_idx].strip().startswith("|"):
        row_line = lines[row_idx]
        cells = [c.strip() for c in row_line.split("|")[1:-1]]
        row = table.add_row()
        for i, cell_text in enumerate(cells):
            if i < len(headers):
                cell = row.cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                # Handle inline formatting in cell
                add_formatted_text(p, cell_text, Pt(10))
        row_idx += 1

    doc.add_paragraph()  # spacing after table
    return row_idx  # next line index


def add_formatted_text(paragraph, text, default_size=Pt(11)):
    """Add text with bold/italic/code parsing to a paragraph."""
    # Bold: **text**
    # Italic: *text*
    # Inline code: `text`
    # Bold+Italic: ***text***

    pattern = r"(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)"
    parts = re.split(pattern, text)

    for part in parts:
        if part is None or part == "":
            continue
        # Check which group matched
        if part.startswith("***") and part.endswith("***"):
            inner = part[3:-3]
            run = paragraph.add_run(inner)
            run.bold = True
            run.italic = True
            run.font.size = default_size
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        elif part.startswith("**") and part.endswith("**"):
            inner = part[2:-2]
            run = paragraph.add_run(inner)
            run.bold = True
            run.font.size = default_size
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            inner = part[1:-1]
            run = paragraph.add_run(inner)
            run.italic = True
            run.font.size = default_size
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        elif part.startswith("`") and part.endswith("`"):
            inner = part[1:-1]
            run = paragraph.add_run(inner)
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xCC, 0x33, 0x33)
        else:
            run = paragraph.add_run(part)
            run.font.size = default_size
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def convert_md_to_docx(md_path, docx_path):
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        # Skip empty lines
        if not line.strip():
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "999999")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Table detection
        if line.strip().startswith("|") and i + 1 < len(lines) and re.match(r"^\|[\s\-:|]+\|$", lines[i + 1].strip()):
            i = add_table(doc, lines, i)
            continue

        # Headings
        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=3)
            for run in p.runs:
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            i += 1
            continue
        elif line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=2)
            for run in p.runs:
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            i += 1
            continue
        elif line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            for run in p.runs:
                run.font.name = "微软雅黑"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            i += 1
            continue

        # Unordered list
        if re.match(r"^[\-\*]\s+", line):
            p = doc.add_paragraph(style="List Bullet")
            add_formatted_text(p, line[2:].strip())
            i += 1
            continue

        # Ordered list
        ol_match = re.match(r"^(\d+)[\.\)]\s+(.+)", line)
        if ol_match:
            p = doc.add_paragraph(style="List Number")
            add_formatted_text(p, ol_match.group(2).strip())
            i += 1
            continue

        # Checkbox list
        cb_match = re.match(r"^\[([ xX])\]\s+(.+)", line)
        if cb_match:
            prefix = "☑ " if cb_match.group(1).lower() == "x" else "☐ "
            p = doc.add_paragraph()
            run = p.add_run(prefix + cb_match.group(2).strip())
            run.font.size = Pt(11)
            run.font.name = "微软雅黑"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            i += 1
            continue

        # Code block
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].rstrip().startswith("```"):
                code_lines.append(lines[i].rstrip())
                i += 1
            i += 1  # skip closing ```

            for cl in code_lines:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run(cl)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
            doc.add_paragraph()  # trailing space
            continue

        # Regular paragraph with possible inline formatting
        p = doc.add_paragraph()
        add_formatted_text(p, line.strip())

        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python md2docx.py <input.md> <output.docx>")
        sys.exit(1)

    md_path = sys.argv[1]
    docx_path = sys.argv[2]

    if not Path(md_path).exists():
        print(f"Error: file not found: {md_path}")
        sys.exit(1)

    convert_md_to_docx(md_path, docx_path)
